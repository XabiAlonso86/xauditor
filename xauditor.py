#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Imports del sistema
import os
import sys
import subprocess
import multiprocessing as mp
import logging
import time
import colorlog
from docx import Document
# Imports de paquetes propios
import sistema.operaciones as oper
import classes.scaner as scn
import classes.pythondocxhandler as pdh


# Variables  globales
analyzedIPs = {} # Diccionario con todas las IPs que se analizan

# Comprobación de parámetros
print ("xauditor alpha 0.3")
if len (sys.argv) < 2:
    print ("<ERROR>: Hay que introducir una dirección IP al menos")
    sys.exit()

# Función para multiproceso
def multProc(targetin, scanip, port):
    jobs = []
    p = mp.Process(target=targetin, args=(scanip,port))
    jobs.append(p)
    p.start()
    return

# Método para crear el log de xauditor
def createLogger(mainPath):

    # Creamos el logger
    logger = colorlog.getLogger("xauditor")
    # Creamos formato de log consola
    formatter = colorlog.ColoredFormatter('%(log_color)s%(levelname)s - %(message)s')
    # Establecemos el nivel
    logger.setLevel(logging.DEBUG)
    # Creamos console handle para mostrar log en la consola
    ch = colorlog.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(formatter)

    # Creamos Formatter para file handler
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    # Creamos file handler para guardar el log en un fichero
    fh = logging.FileHandler(mainPath + '/xauditor@%s.log' % time.strftime("%H_%M_%S"))
    fh.setFormatter(formatter)
    fh.setLevel(logging.DEBUG)

    # Añadir al logger los handlers para que se encargue de generar los logs
    logger.addHandler(fh)
    logger.addHandler(ch)
    return logger

def analyzeIP(ip,folder):
    # Variables para guardar procesos creados para la ip
    jobs = []
    # Manager para compartir diccionarios en los procesos
    manager = mp.Manager()
    servicios = manager.dict()
    serviciosUDP = manager.dict()

    # Creamos documento word para la IP
    document = Document()
    document.add_heading(ip + '_Template', 0)
    # Añadimos escaneos NMAP, conexión UDP, etc
    pdh.addRecon(document,ip,folder)
    # QUITAR CUANDO SE QUIERA PROBAR
    test = False
    if (test == False):
        # Creamos procesos escaneo NMAP
        p = mp.Process(target=scn.nmapScan,name="nmapScan_" + ip, args=(scanip,ipPath,servicios))
        jobs.append(p)
        p.start()
        # UDP Nmap Scan
        p = mp.Process(target=scn.nmapUdpScan,name="nmapUdpScan_" + ip, args=(scanip,ipPath,serviciosUDP))
        jobs.append(p)
        p.start()
        # UnicornScan (sólo guardamos resultado)
        p = mp.Process(target=scn.unicornScan,name="unicornScan_" + ip, args=(scanip,ipPath))
        jobs.append(p)
        p.start()

    # Entramos en bucle hasta que se completen todos los procesos
    # Es posible que los resultados de un proceso impliquen la creación de otro
    print ("Procesos: %s" % (len(jobs)))
    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso %s ha acabado!" % (job.name))
                jobFinished = job.name
                jobs.remove(job)
                # Miramos qué proceso ha terminado y si hay que crear alguno más
                if "nmapScan" in jobFinished:
                    # Añadimos cabecera para resultados
                    document.add_heading('Resultados NMAP Scan %s' % (ip), level=1)
                    checkJobIPFinished(jobFinished,jobs,ip,ipPath,servicios,document)
                elif "nmapUdpScan" in jobFinished:
                    # Añadimos cabecera para resultados
                    document.add_heading('Resultados NMAP UDP Scan %s' % (ip), level=1)
                    checkJobIPFinished(jobFinished,jobs,ip,ipPath,serviciosUDP,document)
                else:
                    checkJobIPFinished(jobFinished,jobs,ip,ipPath)
                logger.debug("Quedan %s" % (len(jobs)))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.info("analyzeIP de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("analyzeIP de IP %s terminado en %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    document.save(folder + '_' + ip + '_template.doc')

    
# Método para comprobar que proceso de análisis de IP ha terminado
# * args recibirá servicios ocasionalmente por eso no se pasa siempre
def checkJobIPFinished(nombre,listaJobs,ip,folder,*args):

    logger.debug("Entro checkJobIPFinished (%s)." % (nombre))
    if "nmapScan" in nombre or "nmapUdpScan" in nombre:
        # Analizamos los servicios encontrados        
        servicios = args[0]
        # Recogemos documento
        document = args[1]
        # Añadimos al documento los servicios encontrados
        pdh.addNmapResults(document,servicios)
        # Lo primero es ver si tenemos más de un serivio (servicio OS siempre está)
        if len(servicios) == 1:
            logger.warning("No se han encontrado servicios para la IP %s. Revise los análisis de NMAP para mayor información" % (ip))
        else:
            for nombre, servicio in servicios.items():
                # Servicio HTTP
                if (nombre == "http") or (nombre == "http-proxy") or (nombre == "http-alt") or (nombre == "http?"):
                    for puerto in servicio.puertos:
                        pdh.addHttp(document,ip,folder,puerto)
                        # Lanzamos proceso para analizar http
                        logger.info("Encontrado servicio http para ip %s en puerto %s" % (ip,puerto))
                        # Creamos proceso y lo añadimos a la lista
                        p = mp.Process(target=httpenum,name="httpenum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso httpenum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio HTTPS
                elif (nombre == "ssl/http") or (nombre == "https") or (nombre == "https?"):
                    for puerto in servicio.puertos:
                        pdh.addHttps(document,ip,folder,puerto)
                        # Lanzamos proceso para analizar https
                        logger.info("Encontrado servicio https para ip %s en puerto %s" % (ip,puerto))
                        # Creamos proceso y lo añadimos a la lista
                        p = mp.Process(target=httpsenum,name="httpsenum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso httpsenum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio SMTP
                elif (nombre == "smtp"):
                    pdh.addSmtp(document,ip,folder)
                    for puerto in servicio.puertos:
                        # Lanzamos proceso para analizar smtp
                        logger.info("Encontrado servicio smtp para ip %s en puerto %s" % (ip,puerto))
                        # Creamos proceso y lo añadimos a la lista
                        p = mp.Process(target=smtpenum,name="smtpenum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso smtpenum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio FTP
                elif (nombre == "ftp"):
                    pdh.addFtp(document,ip,folder)
                    for puerto in servicio.puertos:
                        # Lanzamos proceso para analizar ftp
                        logger.info("Encontrado servicio ftp para ip %s en puerto %s" % (ip,puerto))
                        # Creamos proceso y lo añadimos a la lista
                        p = mp.Process(target=ftpenum,name="ftpenum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso ftpenum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio NetBios
                elif (nombre == "microsoft-ds") or (nombre == "netbios-ssn"):
                    pdh.addSmb(document,ip,folder)
                    for puerto in servicio.puertos:
                        # Lanzamos proceso para analizar netbios
                        logger.info("Encontrado servicio NetBios para ip %s en puerto %s" % (ip,puerto))
                        # Creamos el proceso y lo añadimos a la lista
                        p = mp.Process(target=smbEnum,name="smbEnum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso smbEnum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio ms-sql
                elif (nombre == "ms-sql"):
                    pdh.addMssql(document,ip,folder)
                    for puerto in servicio.puertos:
                        # Lanzamos proceso para analizar ms-sql
                        logger.info("Encontrado servicio MS-SQL para ip %s en puerto %s" % (ip,puerto))
                        # Creamos el proceso y lo añadimos a la lista
                        p = mp.Process(target=mssqlEnum,name="mssqlEnum_" + ip, args=(ip,puerto,folder))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso mssqlEnum a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio ssh
                elif (nombre == "ssh"):
                    for puerto in servicio.puertos:
                        # Lanzamos proceso para analizar ms-sql
                        logger.info("Encontrado servicio ssh para ip %s en puerto %s" % (ip,puerto))
                        # Creamos el proceso y lo añadimos a la lista
                        p = mp.Process(target=scn.conectarPuerto,name="sshScan_" + ip, args=(ip,puerto,folder + "/ssh""ssh"))
                        listaJobs.append(p)
                        logger.debug("Añadido proceso sshScan a la lista de la IP %s. Ahora hay %s" % (ip,len(listaJobs)))
                        # Iniciamos el proceso
                        p.start()
                # Servicio nfs
                elif (nombre == "nfs"):
                    pdh.addNfs(document,ip,folder)
                # Servicio MSRPC
                elif (nombre == "msrpc"):
                    pdh.addMsRPC(document,ip,folder)
                # Servicio SNMP
                elif (nombre == "snmp"):
                    pdh.addSnmp(document,ip,folder)
                # Servicio Oracle
                elif (nombre == "oracle"):
                    pdh.addOracle(document,ip,folder)
                # Servicio Mysql
                elif (nombre == "mysql"):
                    pdh.addMysql(document,ip,folder)
                # Servicio RpcBind
                elif (nombre == "rpcbind"):
                    pdh.addRpcbind(document,ip,folder)
                # Servicio Remote Desktop
                elif (nombre == "ms-wbt-server"):
                    pdh.addRemoteDesktop(document,ip,folder)
                # Servicio Oracle Web UI
                #elif (nombre == "oracle"):
                #    pdh.addOracleWebInterface(document,ip,folder)
                ## Resto de servicios
                            

    logger.debug("Fin checkJobIPFinished.")
    return


# Método para lanzar Nikto, dirb, curl y script de NMAP con servicio HTTP
def httpenum(ip,puerto,folder):
    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/http"
    logger.info ("Iniciando enumeración HTTP para la IP %s:%s" % (ip,puerto))
    # Creamos directorios de las carpetas
    oper.createFolder(path)
    # Creamos procesos
    # dirbScan
    p = mp.Process(target=scn.dirbScan,name="dirbScan_" + ip, args=(ip,puerto,"http",path))
    jobs.append(p)
    p.start()
    # nikto
    p = mp.Process(target=scn.niktoScan,name="niktoScan_" + ip, args=(ip,puerto,"http",path))
    jobs.append(p)
    p.start()
    scn.curlScan(ip,puerto,path)
    p = mp.Process(target=scn.nmapScanHTTP,name="nmapScanHTTP_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()
    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso httpenum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("%s proceso/s restante/s httpenum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("httpenum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin httpenum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return

# Método para lanzar Nikto, dirb, curl y script de NMAP con servicio HTTPS
def httpsenum(ip,puerto,folder):
    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/https"
    logger.info ("Iniciando enumeración HTTPS para la IP %s:%s" % (ip,puerto))
    # Creamos directorios de las carpetas
    oper.createFolder(path)
    # dirbScan
    p = mp.Process(target=scn.dirbScan,name="dirbScanHTTPS_" + ip, args=(ip,puerto,"https",path))
    jobs.append(p)
    p.start()
    # nikto
    p = mp.Process(target=scn.niktoScan,name="niktoScanHTTPS_" + ip, args=(ip,puerto,"https",path))
    jobs.append(p)
    p.start()
    # sslScan
    p = mp.Process(target=scn.sslScan,name="sslScan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()
    #nmapHTTP_process = multiprocessing.Process(target=nmapScanHTTP, args=(ip,puerto,path,"HTTP"))
    #nmapHTTP_process.start()
    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso httpsenum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("Quedan %s procesos httpsenum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("httpsenum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin httpsenum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return

# Método para lanzar banner grabbing y NMAP con servicio SMTP
def smtpenum(ip,puerto,folder):

    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/smtp"
    logger.info ("Iniciando smtpenum para la IP %s:%s" % (ip,puerto))

    # Creamos directorios de las carpetas
    oper.createFolder(path)
    # Banner Grabbing
    p = mp.Process(target=scn.conectarPuerto,name="conectarPuerto_" + ip, args=(ip,puerto,path,"smtp"))
    jobs.append(p)
    p.start()
    # smtpScan
    p = mp.Process(target=scn.smtpScan,name="smtpCan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()

    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso smtpenum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("Quedan %s procesos smtpenum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("smtpenum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin smtpenum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return

# Método para lanzar banner grabbing y NMAP con servicio SMTP
def ftpenum(ip,puerto,folder):

    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/ftp"
    logger.info ("Iniciando ftpenum para la IP %s:%s" % (ip,puerto))

    # Creamos directorios de las carpetas
    oper.createFolder(path)

    # Banner Grabbing
    p = mp.Process(target=scn.conectarPuerto,name="conectarPuerto_" + ip, args=(ip,puerto,path,"ftp"))
    jobs.append(p)
    p.start()
    # smtpScan
    p = mp.Process(target=scn.ftpScan,name="ftpScan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()

    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso ftpenum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("Quedan %s procesos ftpenum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("ftpenum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin ftpenum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return


# Método para lanzar banner grabbing y NMAP con servicio SMTP
def smbEnum(ip,puerto,folder):

    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/smb"
    logger.info ("Iniciando smbEnum para la IP %s:%s" % (ip,puerto))

    # Creamos directorios de las carpetas
    oper.createFolder(path)
    
    # smbNmap
    p = mp.Process(target=scn.smbNmapScan,name="smbNmapScan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()

    # enum4linux
    p = mp.Process(target=scn.enum4linuxScan,name="enum4linuxScan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()

    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso smbEnum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("Quedan %s procesos smbEnum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("smbEnum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin smbEnum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return

# Método para lanzar banner grabbing y NMAP con servicio SMTP
def mssqlEnum(ip,puerto,folder):

    # Variables para guardar procesos creados para la ip
    jobs = []
    path =  folder + "/mssql"
    logger.info ("Iniciando mssqlEnum para la IP %s:%s" % (ip,puerto))

    # Creamos directorios de las carpetas
    oper.createFolder(path)
    
    # mssqlScan
    p = mp.Process(target=scn.mssqlScan,name="mssqlScan_" + ip, args=(ip,puerto,path))
    jobs.append(p)
    p.start()

    timerElapsed = 0
    while len(jobs) > 0:
        for job in jobs:
            job.join(5)
            timerElapsed += 5
            if not job.exitcode is None:
                logger.info("Proceso mssqlEnum de IP %s %s ha acabado!" % (ip, job.name))
                jobs.remove(job)
                logger.debug("Quedan %s procesos mssqlEnum de IP %s" % (len(jobs),ip))
            else:
                # Cada 30 Segundos mostramos lo que lleva la tarea
                if (timerElapsed % 30) == 0:
                    logger.debug("mssqlEnum de IP %s lleva %s minutos" % (ip,time.strftime("%M:%S", time.gmtime(timerElapsed))))
                    line = ""
                    for job in jobs:
                        line += job.name + " "
                    logger.info("Procesos pendientes: %s" % (line))

    logger.info("Fin mssqlEnum para la IP %s:%s (%s)" % (ip,puerto,time.strftime("%M:%S", time.gmtime(timerElapsed))))
    return


# Comienzo del programa
if __name__=='__main__':

    listaIPs = []
    # Recogemos IPs de los argumentos
    targets = sys.argv
    # Quitamos el primer objeto porque es el nombre del programa
    targets.pop(0)

    # Creamos directorio de la aplicación si es necesario
    mainPath = "%s/xauditor" % (os.environ['HOME'])
    if not os.path.exists(mainPath):
        print ("Creando directorio para aplicación xauditor")
        print("Directorio Aplicación: %s" % (mainPath))
        os.makedirs(mainPath)

    # Creamos el logger de la aplicación
    logger = createLogger(mainPath)

    # Recorremos cada IP obtenida y vamos ejecutando análisis de reco
    for scanip in targets:
        if oper.validIP(scanip) is not None:
            # Creamos directorios para la IP
            ipPath = oper.createIPPath(mainPath,scanip)
            # Comenzamos proceso para analizar la IP
            p = mp.Process(target=analyzeIP,name="analyzeIP_" + scanip, args=(scanip,ipPath))
            listaIPs.append(p)
            p.start()
            logger.debug("Fin iteración para ip %s." % (scanip))
        else:
            logger.error ("IP %s tiene un formato incorrecto",scanip)
    
    os.system('stty sane')
